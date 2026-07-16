"""Meeting-document export with a separate frozen sensitive-policy target."""

from __future__ import annotations

from io import BytesIO
import mimetypes
from pathlib import Path
from typing import Any

from docx import Document

from app.sensitive_policy import apply_sensitive_policy


def build_meeting_docx(
    meeting: dict[str, Any],
    export_kind: str = "all",
    *,
    sensitive_rules: list[dict[str, Any]] | None = None,
    include_policy_audit: bool = False,
) -> bytes | tuple[bytes, dict[str, Any]]:
    """Generate a DOCX document and optionally return export-policy audit evidence.

    ``include_policy_audit`` is additive so existing direct callers retain the historical ``bytes``
    result.  Route consumers opt in to persist the rule version and hit metadata used for export.
    """

    lines, audit = _meeting_export_lines(meeting, export_kind, sensitive_rules or [])
    doc = Document()
    doc.add_heading(meeting.get("meetingName", "会议文稿"), 0)
    doc.add_paragraph(f"导出类型：{export_kind}")
    for kind, text in lines:
        if kind == "heading":
            doc.add_heading(text, level=1)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    stream = BytesIO()
    doc.save(stream)
    data = stream.getvalue()
    return (data, audit) if include_policy_audit else data


def build_meeting_text(
    meeting: dict[str, Any],
    export_kind: str = "all",
    *,
    sensitive_rules: list[dict[str, Any]] | None = None,
    include_policy_audit: bool = False,
) -> str | tuple[str, dict[str, Any]]:
    """Generate text export using the exact same frozen export policy as the DOCX path."""

    lines, audit = _meeting_export_lines(meeting, export_kind, sensitive_rules or [])
    text = "\n".join(item for _kind, item in lines)
    return (text, audit) if include_policy_audit else text


def build_word_list_docx(title: str, words: list[str]) -> bytes:
    """把配置词表导出为真正的 DOCX，而不是仅修改文本文件扩展名。"""

    document = Document()
    document.add_heading(title, 0)
    for word in words:
        normalized = str(word or "").strip()
        if normalized:
            document.add_paragraph(normalized, style="List Bullet")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _meeting_export_lines(
    meeting: dict[str, Any], export_kind: str, rules: list[dict[str, Any]]
) -> tuple[list[tuple[str, str]], dict[str, Any]]:
    """Render export-only text and collect one immutable policy audit across every field.

    Stored ``rawText`` and ``text`` are read but never changed here.  Keeping the masking at this
    boundary prevents export rules from becoming recognition replacements or editor mutations.
    """

    lines: list[tuple[str, str]] = []
    hits: list[dict[str, Any]] = []
    rule_version = apply_sensitive_policy("", rules, "export").rule_version

    def masked(value: Any, field: str) -> str:
        nonlocal rule_version
        result = apply_sensitive_policy(str(value or ""), rules, "export")
        rule_version = result.rule_version
        hits.extend(dict(hit, field=field) for hit in result.hits)
        return result.text

    if export_kind in {"all", "transcript"}:
        lines.append(("heading", "会议转写"))
        for segment in meeting.get("segments", []):
            start = _format_ms(segment.get("startMs", 0))
            speaker = segment.get("speakerName", "发言人")
            lines.append(("text", f"[{start}] {speaker}：{masked(segment.get('text'), 'segments.text')}"))
    summary = meeting.get("summary")
    if isinstance(summary, dict) and export_kind in {"all", "summary"}:
        lines.extend(
            [
                ("heading", "AI 摘要"),
                ("text", f"主题：{masked(summary.get('topic'), 'summary.topic')}"),
                ("text", f"概要：{masked(summary.get('overview'), 'summary.overview')}"),
                ("heading", "会议要点"),
            ]
        )
        for index, point in enumerate(summary.get("keyPoints", [])):
            lines.append(("bullet", masked(point, f"summary.keyPoints[{index}]")))
        lines.append(("heading", "待办事项"))
        for index, todo in enumerate(summary.get("todos", [])):
            if isinstance(todo, dict):
                title = masked(todo.get("title"), f"summary.todos[{index}].title")
                content = masked(todo.get("content"), f"summary.todos[{index}].content")
                lines.append(("bullet", f"{title}：{content}"))
    minutes = meeting.get("minutes")
    if isinstance(minutes, dict) and export_kind in {"all", "minutes"}:
        lines.append(("heading", "会议纪要"))
        for index, paragraph in enumerate(str(minutes.get("content") or "").splitlines()):
            if paragraph.strip():
                lines.append(("text", masked(paragraph.strip(), f"minutes.content[{index}]")))
    return lines, {"target": "export", "ruleVersion": rule_version, "hits": hits}


def read_audio_for_playback(file_record: dict[str, Any]) -> tuple[bytes, str, str]:
    """Return original recorded-media bytes with a truthful filename and MIME type.

    Renaming arbitrary WAV/M4A/MP4 bytes to ``.mp3`` makes browsers select the wrong decoder and
    breaks imported-record playback.  Until a real ffmpeg MP3 transcode is introduced, preserving
    the source container is both lossless and standards-compliant.
    """

    path = Path(file_record["path"])
    data = path.read_bytes()
    output_name = str(file_record.get("filename") or file_record.get("fileName") or path.name)
    guessed_type = mimetypes.guess_type(output_name)[0]
    media_type = str(file_record.get("contentType") or guessed_type or "application/octet-stream")
    return data, output_name, media_type


def _format_ms(value: int) -> str:
    """Format milliseconds as a stable mm:ss transcript label."""

    seconds = int(value) // 1000
    return f"{seconds // 60:02d}:{seconds % 60:02d}"
